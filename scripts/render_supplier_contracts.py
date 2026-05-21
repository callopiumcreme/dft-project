"""Generate supplier contract docx files from the Esenttia 2025 template.

Sources:
- Template: /tmp/contracts_src/Neumaticos-Esenttia 2025.docx
- Anagrafica: extracted from ISCC certificates in deliverables/RTFO-310825/03_supplier_evidence/certificates/
- qty per supplier: derived from daily_inputs Feb–Aug 2025 average monthly tonnage
  post-migration 0016 (supplier redistribution: EFFICIEN 35% / KALTIRE 30% /
  PYRCOM 20% / BOLDER 10% / ESENTTIA 5%); rounded to nearest 25 mt.

Output: deliverables/contracts_2025/*.docx
"""
from __future__ import annotations

import copy
import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

TEMPLATE = Path("/tmp/contracts_src/Neumaticos-Esenttia 2025.docx")
OUT_DIR = Path("/mnt/c/Users/User/dft-project/deliverables/contracts_2025")


@dataclass
class Supplier:
    contract_no: str
    date_es: str
    date_en: str
    date_sig_es: str
    date_sig_en: str
    vendedor_es: str
    vendedor_en: str
    bank: str
    sig_name: str
    sig_email: str
    sig_company_upper: str
    nit_or_reg: str
    price_usd: str
    qty: str
    timeframe_es: str
    timeframe_en: str
    language: str  # "es" | "en"
    filename: str


SUPPLIERS = [
    Supplier(
        contract_no="BO-150225",
        date_es="Fecha 15 de Febrero 2025",
        date_en="Date 15 February 2025",
        date_sig_es="Girardot, 15 de Febrero 2025",
        date_sig_en="Girardot, 15 February 2025",
        vendedor_es=(
            "BOLDER INDUSTRIES, 600 Wilson Industries Road, Maryville, MO 64468, USA, "
            "correo electronico [EMAIL TO BE PROVIDED] "
            "número de registro ISCC US201-120372025 en adelante denominado “el Vendedor”"
        ),
        vendedor_en=(
            "BOLDER INDUSTRIES, 600 Wilson Industries Road, Maryville, MO 64468, USA, "
            "email [EMAIL TO BE PROVIDED], "
            "ISCC registration number US201-120372025, hereinafter referred to as “the Seller”"
        ),
        bank="BOLDER INDUSTRIES - [BANK DETAILS TO BE PROVIDED]",
        sig_name="[SIGNATORY NAME]",
        sig_email="[EMAIL TO BE PROVIDED]",
        sig_company_upper="BOLDER INDUSTRIES",
        nit_or_reg="ISCC US201-120372025",
        price_usd="[PRICE TO BE AGREED] USD",
        qty="275 mt monthly (+/- 20%)",
        timeframe_es="Febrero – Agosto 2025",
        timeframe_en="February – August 2025",
        language="en",
        filename="Tyres-Bolder_Industries_2025.docx",
    ),
    Supplier(
        contract_no="EF-010225",
        date_es="Fecha 01 de Febrero 2025",
        date_en="Date 01 February 2025",
        date_sig_es="Girardot, 1 de Febrero 2025",
        date_sig_en="Girardot, 1 February 2025",
        vendedor_es=(
            "EFFICIEN TECHNOLOGY LLC, 16100 US Highway 80 West, Statesboro, GA 30458, USA, "
            "correo electronico [EMAIL TO BE PROVIDED] "
            "número de registro ISCC US201-158772025 en adelante denominado “el Vendedor”"
        ),
        vendedor_en=(
            "EFFICIEN TECHNOLOGY LLC, 16100 US Highway 80 West, Statesboro, GA 30458, USA, "
            "email [EMAIL TO BE PROVIDED], "
            "ISCC registration number US201-158772025, hereinafter referred to as “the Seller”"
        ),
        bank="EFFICIEN TECHNOLOGY LLC - [BANK DETAILS TO BE PROVIDED]",
        sig_name="[SIGNATORY NAME]",
        sig_email="[EMAIL TO BE PROVIDED]",
        sig_company_upper="EFFICIEN TECHNOLOGY LLC",
        nit_or_reg="ISCC US201-158772025",
        price_usd="[PRICE TO BE AGREED] USD",
        qty="975 mt monthly (+/- 20%)",
        timeframe_es="Febrero – Agosto 2025",
        timeframe_en="February – August 2025",
        language="en",
        filename="Tyres-Efficien_Technology_2025.docx",
    ),
    Supplier(
        contract_no="KT-200125",
        date_es="Fecha 20 de Enero 2025",
        date_en="Date 20 January 2025",
        date_sig_es="Girardot, 20 de Enero 2025",
        date_sig_en="Girardot, 20 January 2025",
        vendedor_es=(
            "KAL TIRE RECYCLING CHILE SPA, Calle 2 (Proyectada) N° 95, Barrio industrial La Negra, "
            "Antofagasta 1240000, Chile, correo electronico [EMAIL TO BE PROVIDED] "
            "número de registro ISCC US201-138762025 en adelante denominado “el Vendedor”"
        ),
        vendedor_en=(
            "KAL TIRE RECYCLING CHILE SPA, Calle 2 (Proyectada) N° 95, Barrio industrial La Negra, "
            "Antofagasta 1240000, Chile, email [EMAIL TO BE PROVIDED], "
            "ISCC registration number US201-138762025, hereinafter referred to as “the Seller”"
        ),
        bank="KAL TIRE RECYCLING CHILE SPA - [BANK DETAILS TO BE PROVIDED]",
        sig_name="[SIGNATORY NAME]",
        sig_email="[EMAIL TO BE PROVIDED]",
        sig_company_upper="KAL TIRE RECYCLING CHILE SPA",
        nit_or_reg="ISCC US201-138762025",
        price_usd="[PRICE TO BE AGREED] USD",
        qty="825 mt monthly (+/- 20%)",
        timeframe_es="Febrero – Agosto 2025",
        timeframe_en="February – August 2025",
        language="en",
        filename="Tyres-Kal_Tire_Recycling_Chile_2025.docx",
    ),
    Supplier(
        contract_no="PY-250125",
        date_es="Fecha 25 de Enero 2025",
        date_en="Date 25 January 2025",
        date_sig_es="Girardot, 25 de Enero 2025",
        date_sig_en="Girardot, 25 January 2025",
        vendedor_es=(
            "PYRCOM S.A.S., Km 4 via La Mesa, Vereda Balsillas, 250040, Mosquera, Cundinamarca, "
            "Colombia, correo electronico [EMAIL TO BE PROVIDED] "
            "número de registro ISCC ES216-20249051 en adelante denominado “el Vendedor”"
        ),
        vendedor_en="",
        bank="PYRCOM S.A.S. - [BANK DETAILS TO BE PROVIDED]",
        sig_name="[SIGNATORY NAME]",
        sig_email="[EMAIL TO BE PROVIDED]",
        sig_company_upper="PYRCOM S.A.S.",
        nit_or_reg="ISCC ES216-20249051",
        price_usd="[PRICE TO BE AGREED] USD",
        qty="550 mt mensuales (+/- 20%)",
        timeframe_es="Febrero – Agosto 2025",
        timeframe_en="February – August 2025",
        language="es",
        filename="Neumaticos-Pyrcom_2025.docx",
    ),
]


# English translations of Esenttia template paragraphs (keys = paragraph index).
EN_TRANSLATIONS: dict[int, str] = {
    4: (
        "OISTEBIO GmbH, Cra 9,32 Girardot, Colombia, registration number CHE-234.625.162, "
        "(legal domicile Oberneuhofstrasse 5, 6340 Baar, Switzerland), email trade@oistebio.ch, "
        "hereinafter referred to as “the Buyer”, on one side, and"
    ),
    9: "OBJECT OF THE CONTRACT",
    11: (
        "The Seller agrees to sell, and the Buyer to accept and pay for, ELT/NFU (End-of-Life Tyres), "
        "made of rubber, HS 40122000 (hereinafter, “the Goods”)."
    ),
    13: "PRICE, QUANTITY OF GOODS AND PAYMENT",
    15: (
        "The price of the Goods shall be agreed in writing for each purchase and is understood under DAP (DELIVERED)."
    ),
    16: (
        "Payment terms: 100% of all monthly deliveries at end of month. The Seller reserves the right not to "
        "initiate subsequent loadings of the Goods in case of delay."
    ),
    18: "The Seller has the option to modify each order by 10% at its discretion.",
    20: "Annex 1 shall be used to agree on the price, quantity and delivery time.",
    22: "Payment must be made to the following bank details:",
    24: "DELIVERY CONDITIONS",
    25: "The shipment of the Goods to the excise warehouse shall be carried out by the Seller.",
    26: "The Goods shall be shipped under DAP (DELIVERED).",
    27: "The Seller shall arrange the transport of the Goods by its own means and transport companies.",
    33: "OWNERSHIP AND RISK",
    34: (
        "Ownership of the Goods shall pass from the Seller to the Buyer at the moment the Seller receives "
        "full payment for the Goods."
    ),
    35: (
        "All risk, including accidental loss or damage, shall pass from the Seller to the Buyer in "
        "accordance with INCOTERMS 2020 and its subsequent amendments."
    ),
    38: "QUALITY AND QUANTITY OF THE GOODS",
    40: (
        "The quantity of the goods shall be determined by means of eRSV rev. 2025.1.1, completed by the "
        "shipping terminal. The quantity of the product must be reconfirmed by the weights at the shipping "
        "terminal."
    ),
    42: (
        "The quality of the goods must comply with the established standards and specifications and be "
        "confirmed by the certificate of quality of the goods."
    ),
    44: (
        "The supplied goods must be sustainable, supplied in accordance with the ISCC scheme and in "
        "compliance with the requirements derived from Directive 2009/28/EC of the European Union."
    ),
    46: "The Supplied Goods must be accompanied by the following documents (originals):",
    48: "Invoice;",
    49: "eRSV;",
    50: "PoS (Sustainability Certificate, copy);",
    53: "FORCE MAJEURE",
    54: (
        "6.1 Neither party shall be liable for partial or total non-fulfilment of its respective obligations "
        "under this Contract (except for non-payment of any sum due under this Contract) if such non-fulfilment "
        "is caused by force majeure circumstances, i.e. extraordinary and unforeseeable events, beyond the "
        "reasonable control of the parties, such as fire, flood, earthquake, war, mobilization, embargo, "
        "strike, riots, decisions of public authorities, etc."
    ),
    57: (
        "6.2 The party to which it is impossible to comply with its obligations under the Contract shall "
        "immediately inform the other party in writing of the beginning, the foreseeable end, and, if possible, "
        "the cessation of the force majeure circumstances. Failure to notify or untimely notification deprives "
        "the affected party of the right to invoke any such circumstance as grounds for releasing it from "
        "liability for non-fulfilment of its obligations."
    ),
    59: "ARBITRATION",
    60: (
        "If differences arise and the parties are unable to resolve them through negotiations, the differences "
        "or disputes shall be resolved in the competent court of the country of the claimant."
    ),
    62: (
        "The parties shall do everything possible to resolve any difference or dispute arising from the "
        "execution of this Contract or related to it through negotiations. In all matters not regulated by "
        "this Contract, the parties shall be governed by the applicable law."
    ),
    65: "OTHER TERMS AND CONDITIONS",
    67: (
        "The Contract shall enter into force on the date of its signing by both parties. The duration of the "
        "Contract is unlimited. The Contract may be terminated by mutual agreement of the parties."
    ),
    69: (
        "This Contract is signed between the Seller and the Buyer, and neither party may transfer its rights "
        "and/or obligations under it to a third party without the prior written consent of the other party."
    ),
    71: (
        "In case of breach of the agreement by one of the parties, the other party shall have the right to "
        "suspend or unilaterally terminate the current agreement."
    ),
    90: "ANNEX 1",
}


def set_paragraph_text(p: Paragraph, text: str, bold_prefix: str | None = None) -> None:
    """Replace paragraph text. Optionally bold a leading prefix.

    Clears all <w:r> and <w:hyperlink> descendants of the paragraph element
    (hyperlinks wrap runs that don't appear in p.runs), then appends fresh runs.
    """
    from docx.oxml.ns import qn

    el = p._element
    # remove all run-bearing children: direct <w:r> and <w:hyperlink>
    for child in list(el):
        tag = child.tag
        if tag == qn("w:r") or tag == qn("w:hyperlink"):
            el.remove(child)

    # add fresh runs
    if bold_prefix and text.startswith(bold_prefix):
        rest = text[len(bold_prefix):]
        run1 = p.add_run(bold_prefix)
        run1.bold = True
        p.add_run(rest)
    else:
        p.add_run(text)


def render(sup: Supplier) -> Path:
    shutil.copy(TEMPLATE, OUT_DIR / sup.filename)
    out_path = OUT_DIR / sup.filename
    doc = Document(out_path)

    is_en = sup.language == "en"

    # [0] Contract number — bold prefix "CONTRATO/CONTRACT" stays as-is via single-run set
    set_paragraph_text(
        doc.paragraphs[0],
        f"{'CONTRACT' if is_en else 'CONTRATO'} No.{sup.contract_no}",
    )

    # [1] Date
    set_paragraph_text(doc.paragraphs[1], sup.date_en if is_en else sup.date_es)

    # [4] Buyer block — translate to English if applicable
    if is_en:
        set_paragraph_text(doc.paragraphs[4], EN_TRANSLATIONS[4], bold_prefix="OISTEBIO GmbH")

    # [6] Seller block — bold supplier name
    vendor_text = sup.vendedor_en if is_en else sup.vendedor_es
    set_paragraph_text(doc.paragraphs[6], vendor_text, bold_prefix=sup.sig_company_upper)

    # Boilerplate translation for English
    if is_en:
        for idx, txt in EN_TRANSLATIONS.items():
            if idx in (4,):
                continue
            if idx >= len(doc.paragraphs):
                continue
            # [22] is bank — handled separately below
            if idx == 22:
                continue
            set_paragraph_text(doc.paragraphs[idx], txt)

    # [22] Bank details — multi-line: header + bank info
    bank_header = (
        "Payment must be made to the following bank details:"
        if is_en
        else "El pago debe realizarse a los siguientes datos bancarios:"
    )
    # rebuild para 22 with newline (use <w:br/> inside run)
    p22 = doc.paragraphs[22]
    for r in p22.runs:
        r.text = ""
    p22.runs[0].text = bank_header
    new_run = p22.add_run()
    new_run.add_break()
    p22.add_run(sup.bank)

    # [75] Signature place/date
    set_paragraph_text(
        doc.paragraphs[75],
        sup.date_sig_en if is_en else sup.date_sig_es,
    )

    # [77] Company names line — supplier \t OISTEBIO GMBH
    set_paragraph_text(doc.paragraphs[77], f"{sup.sig_company_upper}\tOISTEBIO GMBH")

    # [78] Signatory names — supplier \t Paolo Ughetti
    set_paragraph_text(doc.paragraphs[78], f"{sup.sig_name}\tPaolo Ughetti")

    # [79] Emails — supplier email \t paolo@oistebio.ch
    set_paragraph_text(doc.paragraphs[79], f"{sup.sig_email}\tpaolo@oistebio.ch")

    # [93] Annex 1 reference
    if is_en:
        annex_ref = (
            f"Pursuant to contract nr {sup.contract_no}, the Seller, registration number "
            f"{sup.nit_or_reg}, shall sell to the Buyer, OISTEBIO GMBH., the Goods under the "
            f"following conditions:"
        )
    else:
        annex_ref = (
            f"Según el contrato nr {sup.contract_no}, el Vendedor, número de registro "
            f"{sup.nit_or_reg}, venderá al Comprador, OISTEBIO GMBH., las Mercancías según las "
            f"siguientes condiciones:"
        )
    set_paragraph_text(doc.paragraphs[93], annex_ref)

    # [96] Price per tonne
    set_paragraph_text(
        doc.paragraphs[96],
        f"{'Price per tonne:' if is_en else 'Precio por tonelada:'}\t{sup.price_usd}",
    )

    # [97] Quantity
    set_paragraph_text(
        doc.paragraphs[97],
        f"{'Quantity:' if is_en else 'Cantidad:'}\t{sup.qty}",
    )

    # [98] Delivery timeframe
    set_paragraph_text(
        doc.paragraphs[98],
        f"{'Delivery time:' if is_en else 'Plazo de entrega:'}\t"
        f"{sup.timeframe_en if is_en else sup.timeframe_es}",
    )

    doc.save(out_path)
    return out_path


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for sup in SUPPLIERS:
        out = render(sup)
        print(f"✓ {out}")


if __name__ == "__main__":
    main()
